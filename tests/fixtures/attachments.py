"""Real attachment bytes for tests — genuinely generated, not stubbed."""

from __future__ import annotations

import base64
import io
import zipfile


def docx_bytes(paragraphs: list[str] | None = None) -> bytes:
    """A real .docx file containing the given paragraphs."""
    import docx

    document = docx.Document()
    for paragraph in paragraphs or ["Hello from a Word document."]:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def docx_with_table(rows: list[list[str]]) -> bytes:
    import docx

    document = docx.Document()
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def pdf_bytes(pages: int = 1) -> bytes:
    """A real (blank) PDF. pypdf can't author text, so this has none."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def pdf_with_text(text: str = "Statement of account") -> bytes:
    """A minimal hand-built PDF whose text pypdf can actually extract."""
    stream = f"BT /F1 12 Tf 20 100 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for index, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n"
    ).encode()
    return bytes(out)


def zip_bomb(entry_size: int = 200 * 1024 * 1024) -> bytes:
    """A small archive whose index claims an enormous uncompressed size.

    Built with a real, highly-compressible payload so the ZIP index reports a
    huge `file_size` from a tiny archive — the shape of an actual bomb.
    """
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        archive.writestr("[Content_Types].xml", b"\x00" * entry_size)
    return buffer.getvalue()


def windows_executable() -> bytes:
    """Bytes that begin with the Windows PE signature."""
    return b"MZ\x90\x00\x03\x00\x00\x00" + b"\x00" * 128


def linux_executable() -> bytes:
    return b"\x7fELF\x02\x01\x01\x00" + b"\x00" * 128


def gmail_attachment_payload(data: bytes) -> dict[str, str]:
    """What the Gmail attachments endpoint returns."""
    return {
        "size": str(len(data)),
        "data": base64.urlsafe_b64encode(data).decode("ascii").rstrip("="),
    }


__all__ = (
    "docx_bytes",
    "docx_with_table",
    "gmail_attachment_payload",
    "linux_executable",
    "pdf_bytes",
    "pdf_with_text",
    "windows_executable",
    "zip_bomb",
)
