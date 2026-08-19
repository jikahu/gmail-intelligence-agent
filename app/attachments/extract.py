"""Reading text out of attachments, safely (CLAUDE.md §11).

The single rule this module exists to uphold: **attachments are read as
information, never executed.**

There is no ``subprocess``, no ``eval``, no ``os.system``, and no path that
hands a file to the operating system to open. The libraries used are parsers:
``pypdf`` reads a PDF's text objects and ``python-docx`` reads the XML inside
the archive. Neither runs the JavaScript a PDF may embed nor the VBA a Word
document may contain — and macro-enabled formats (``.docm``, ``.xlsm``) are
refused before any library sees them.

Every extractor is bounded by the caps in :mod:`app.attachments.types`, and
every one returns a result rather than raising. A file we can't read produces a
recorded failure, which by design changes nothing about how the email is
classified.
"""

from __future__ import annotations

import csv
import io
import logging
import re
import zipfile
from typing import Callable

from app.attachments.models import ExtractedAttachment, ExtractionStatus
from app.attachments.types import (
    DANGEROUS_EXTENSIONS,
    MAX_ATTACHMENT_BYTES,
    MAX_COMPRESSION_RATIO,
    MAX_CSV_ROWS,
    MAX_DOCX_PARAGRAPHS,
    MAX_EXTRACTED_CHARS,
    MAX_PDF_PAGES,
    MAX_UNCOMPRESSED_BYTES,
    AttachmentKind,
    classify_attachment,
    extension_of,
    kind_from_magic,
    safe_display_name,
)
from app.logging_config import get_logger

log = get_logger("app.attachments.extract")

# pypdf narrates malformed files at WARNING ("EOF marker not found" and
# friends). We already report that outcome ourselves as CORRUPTED, so its
# commentary is noise in our logs — and malformed PDFs are routine in email.
logging.getLogger("pypdf").setLevel(logging.ERROR)

_WHITESPACE = re.compile(r"[ \t ]+")
_BLANK_LINES = re.compile(r"\n{3,}")
_CONTROL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _tidy(text: str) -> tuple[str, bool]:
    """Normalize extracted text and truncate. Returns ``(text, truncated)``."""
    if not text:
        return "", False
    cleaned = _CONTROL.sub(" ", text)
    cleaned = _WHITESPACE.sub(" ", cleaned)
    cleaned = _BLANK_LINES.sub("\n\n", cleaned)
    cleaned = cleaned.strip()
    if len(cleaned) > MAX_EXTRACTED_CHARS:
        return cleaned[:MAX_EXTRACTED_CHARS], True
    return cleaned, False


def _decode(data: bytes) -> str:
    """Decode bytes as text, trying the encodings email actually uses."""
    for encoding in ("utf-8", "utf-16", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


# --------------------------------------------------------------------
# Per-kind extractors
# --------------------------------------------------------------------


def _extract_text(data: bytes, result: ExtractedAttachment) -> None:
    text, truncated = _tidy(_decode(data))
    result.text = text
    result.truncated = truncated
    result.status = ExtractionStatus.EXTRACTED if text else ExtractionStatus.EMPTY


def _extract_csv(data: bytes, result: ExtractedAttachment) -> None:
    """Read a CSV as rows of text.

    Note on "CSV injection": a cell beginning ``=``, ``+``, ``-`` or ``@`` is
    dangerous only to a spreadsheet application that evaluates it. Nothing here
    evaluates anything — the cell is text to us, and we never write a CSV back
    out — so no formula can run. Cells are read exactly as they are.
    """
    raw = _decode(data)
    try:
        dialect = csv.Sniffer().sniff(raw[:4096], delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel  # type: ignore[assignment]

    rows: list[str] = []
    truncated = False
    try:
        reader = csv.reader(io.StringIO(raw), dialect)
        for index, row in enumerate(reader):
            if index >= MAX_CSV_ROWS:
                truncated = True
                break
            cells = [str(cell).strip() for cell in row if str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
    except csv.Error as exc:
        result.status = ExtractionStatus.CORRUPTED
        result.error = f"csv parse error: {exc}"
        return

    text, text_truncated = _tidy("\n".join(rows))
    result.text = text
    result.row_count = len(rows)
    result.truncated = truncated or text_truncated
    result.status = ExtractionStatus.EXTRACTED if text else ExtractionStatus.EMPTY


def _extract_pdf(data: bytes, result: ExtractedAttachment) -> None:
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError:
        result.status = ExtractionStatus.LIBRARY_MISSING
        result.error = "pypdf is not installed"
        return

    warnings = list(result.warnings)
    try:
        reader = PdfReader(io.BytesIO(data), strict=False)

        if getattr(reader, "is_encrypted", False):
            # An empty-password PDF is common and harmless; a real one is not
            # something we try to break into.
            try:
                if reader.decrypt("") == 0:
                    result.status = ExtractionStatus.ENCRYPTED
                    result.error = "password protected"
                    return
            except Exception:  # noqa: BLE001
                result.status = ExtractionStatus.ENCRYPTED
                result.error = "password protected"
                return

        # Note, but never run, embedded active content.
        if _pdf_has_active_content(reader):
            warnings.append(
                "this PDF contains embedded script content, which was not run"
            )

        pages = reader.pages
        result.page_count = len(pages)
        chunks: list[str] = []
        truncated = len(pages) > MAX_PDF_PAGES

        for page in list(pages)[:MAX_PDF_PAGES]:
            try:
                chunks.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001 — one bad page isn't fatal
                warnings.append("part of this PDF could not be read")
            if sum(len(chunk) for chunk in chunks) > MAX_EXTRACTED_CHARS:
                truncated = True
                break

        text, text_truncated = _tidy("\n".join(chunks))
        result.text = text
        result.truncated = truncated or text_truncated
        result.warnings = tuple(dict.fromkeys(warnings))
        result.status = (
            ExtractionStatus.EXTRACTED if text else ExtractionStatus.EMPTY
        )
        if not text:
            result.error = "no extractable text (it may be a scan)"
    except PdfReadError as exc:
        result.status = ExtractionStatus.CORRUPTED
        result.error = f"unreadable pdf: {exc}"
    except Exception as exc:  # noqa: BLE001
        result.status = ExtractionStatus.FAILED
        result.error = str(exc)[:200]


def _pdf_has_active_content(reader: object) -> bool:
    """Detect embedded JavaScript or launch actions. Detect only — never run."""
    try:
        root = reader.trailer["/Root"]  # type: ignore[index]
    except Exception:  # noqa: BLE001
        return False
    for key in ("/JavaScript", "/JS", "/OpenAction", "/AA", "/Launch", "/EmbeddedFiles"):
        try:
            if key in root:
                return True
            names = root.get("/Names")
            if names is not None and key in names:
                return True
        except Exception:  # noqa: BLE001
            continue
    return False


def _extract_docx(data: bytes, result: ExtractedAttachment) -> None:
    bomb = _decompression_bomb_reason(data)
    if bomb:
        result.status = ExtractionStatus.TOO_LARGE
        result.error = bomb
        return

    try:
        import docx  # python-docx
    except ImportError:
        result.status = ExtractionStatus.LIBRARY_MISSING
        result.error = "python-docx is not installed"
        return

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 — includes zipfile.BadZipFile
        result.status = ExtractionStatus.CORRUPTED
        result.error = f"unreadable docx: {str(exc)[:150]}"
        return

    chunks: list[str] = []
    truncated = False
    try:
        for index, paragraph in enumerate(document.paragraphs):
            if index >= MAX_DOCX_PARAGRAPHS:
                truncated = True
                break
            if paragraph.text.strip():
                chunks.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))
    except Exception as exc:  # noqa: BLE001
        result.status = ExtractionStatus.FAILED
        result.error = str(exc)[:200]
        return

    text, text_truncated = _tidy("\n".join(chunks))
    result.text = text
    result.truncated = truncated or text_truncated
    result.status = ExtractionStatus.EXTRACTED if text else ExtractionStatus.EMPTY


def _decompression_bomb_reason(data: bytes) -> str | None:
    """Refuse an archive that claims to expand enormously.

    A DOCX is a ZIP. The archive index states each entry's uncompressed size
    *before* anything is decompressed, so this check happens without ever
    expanding the payload — which is the whole point.
    """
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            total = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile:
        return None  # Not a valid zip; the extractor will report it as corrupt.
    except Exception:  # noqa: BLE001
        return None

    if total > MAX_UNCOMPRESSED_BYTES:
        return f"expands to {total // (1024 * 1024)} MB, which is over the limit"
    if data and total / max(len(data), 1) > MAX_COMPRESSION_RATIO:
        return "compression ratio looks like a decompression bomb"
    return None


_EXTRACTORS: dict[AttachmentKind, Callable[[bytes, ExtractedAttachment], None]] = {
    AttachmentKind.TEXT: _extract_text,
    AttachmentKind.CSV: _extract_csv,
    AttachmentKind.PDF: _extract_pdf,
    AttachmentKind.DOCX: _extract_docx,
}


# --------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------


def extract(
    filename: str,
    data: bytes | None,
    mime_type: str = "",
    declared_size: int = 0,
) -> ExtractedAttachment:
    """Read one attachment. Never raises, never executes anything."""
    display_name = safe_display_name(filename)
    size = len(data) if data else declared_size
    kind = classify_attachment(filename, mime_type, data)

    result = ExtractedAttachment(
        filename=display_name,
        mime_type=(mime_type or "").strip().lower(),
        size_bytes=size,
        kind=kind,
    )

    # 1. Executables are refused without being opened.
    if kind is AttachmentKind.EXECUTABLE:
        result.status = ExtractionStatus.BLOCKED
        result.error = "executable or macro-enabled file"
        # No warning here: STATUS_EXPLANATIONS already says it wasn't run, and
        # repeating it made the user-facing sentence say the same thing twice.
        # A warning is for information the status doesn't already carry — such
        # as the disguise below.
        if extension_of(filename) not in DANGEROUS_EXTENSIONS:
            result.warnings = (
                f"the name {display_name!r} does not look like a program, but "
                "the contents are one",
            )
        log.warning(
            "attachment_blocked_executable",
            extra={"attachment_name": display_name, "mime": result.mime_type},
        )
        return result

    # 2. Size gate, before anything is parsed.
    if size > MAX_ATTACHMENT_BYTES:
        result.status = ExtractionStatus.TOO_LARGE
        result.error = f"{size} bytes exceeds the {MAX_ATTACHMENT_BYTES} byte limit"
        return result

    # `None` means the download failed; empty bytes mean the file really is
    # empty. They are different outcomes and get different statuses.
    if data is None:
        result.status = ExtractionStatus.NOT_ATTEMPTED
        result.error = "the attachment could not be downloaded"
        return result

    if not data:
        result.status = ExtractionStatus.EMPTY
        result.error = "the file is empty"
        return result

    # 3. Note when the contents disagree with the name. Not fatal — plenty of
    #    legitimate mail has a wrong MIME type — but worth recording.
    actual = kind_from_magic(data)
    if actual is not None and actual is not kind and kind is not AttachmentKind.UNKNOWN:
        if not (actual is AttachmentKind.ARCHIVE and kind is AttachmentKind.DOCX):
            result.warnings = result.warnings + (
                f"named like a {kind.value} file, but the contents look like "
                f"{actual.value}",
            )

    # 4. Kinds we deliberately don't open.
    extractor = _EXTRACTORS.get(kind)
    if extractor is None:
        result.status = ExtractionStatus.UNSUPPORTED
        result.error = _unsupported_reason(kind)
        return result

    try:
        extractor(data, result)
    except Exception as exc:  # noqa: BLE001 — the last line of defence
        log.warning(
            "attachment_extraction_failed",
            extra={"attachment_name": display_name, "error": str(exc)},
        )
        result.status = ExtractionStatus.FAILED
        result.error = str(exc)[:200]

    return result


def _unsupported_reason(kind: AttachmentKind) -> str:
    if kind is AttachmentKind.IMAGE:
        return "images are not read for text in V1 (no OCR)"
    if kind is AttachmentKind.ARCHIVE:
        return "archives are not opened"
    return f"{kind.value} files are not read for text"


__all__ = ("extract",)
